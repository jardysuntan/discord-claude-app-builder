"""
api.py — FastAPI HTTP server for the discord-claude-app-builder.
Runs alongside the Discord bot as a separate process on port 8100.

Start: python api.py
Or:    uvicorn api:app --host 0.0.0.0 --port 8100
"""

import logging
import os
import secrets
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, HTTPException, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

import service
from agent_factory import get_provider_capabilities
from accounts import AccountManager, Account

# ── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(name)s %(levelname)s %(message)s")
logger = logging.getLogger("api")

# ── Legacy Auth (kept for backward compat) ──────────────────────────────────

TOKEN_FILE = Path(__file__).parent / ".api-token"

def _get_api_token() -> str:
    token = os.getenv("API_TOKEN")
    if token:
        return token
    if TOKEN_FILE.exists():
        return TOKEN_FILE.read_text().strip()
    # Generate and persist a token
    token = secrets.token_urlsafe(32)
    TOKEN_FILE.write_text(token + "\n")
    TOKEN_FILE.chmod(0o600)
    logger.info(f"Generated API token → {TOKEN_FILE}")
    return token

API_TOKEN = _get_api_token()

# ── Account Manager (singleton) ─────────────────────────────────────────────

_account_mgr: AccountManager | None = None

def _get_account_mgr() -> AccountManager:
    global _account_mgr
    if _account_mgr is None:
        _account_mgr = AccountManager()
    return _account_mgr


# ── Multi-tenant Auth ───────────────────────────────────────────────────────

async def get_current_account(authorization: str = Header(...)) -> Account:
    """Authenticate via Bearer token. Supports legacy .api-token and new API keys."""
    if not authorization.startswith("Bearer "):
        raise HTTPException(401, "Missing Bearer token")
    token = authorization[7:]

    # 1. Check legacy .api-token (maps to admin account)
    if token == API_TOKEN:
        mgr = _get_account_mgr()
        # Find admin account (or any account with this legacy key)
        acct = mgr.authenticate(token)
        if acct:
            return acct
        # Fallback: find first admin account
        for a in mgr.list_accounts():
            if a.role == "admin":
                return a
        # No accounts at all — create a synthetic admin for backward compat
        return Account(
            account_id="legacy_admin",
            display_name="Admin (legacy)",
            role="admin",
        )

    # 2. Check new API keys
    mgr = _get_account_mgr()
    acct = mgr.authenticate(token)
    if acct:
        return acct

    raise HTTPException(403, "Invalid token")


# ── App ──────────────────────────────────────────────────────────────────────

app = FastAPI(title="App Builder API", version="1.0.0", docs_url="/api/docs")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("startup")
async def startup():
    service.init()
    logger.info(f"API ready on port {os.getenv('API_PORT', '8100')}")
    logger.info(f"Token: {API_TOKEN[:8]}...")


# ── Request/Response models ──────────────────────────────────────────────────

class BuildAppRequest(BaseModel):
    description: str
    app_name: str | None = None
    platform: str = "web"
    skip_supabase: bool = False
    webhook_url: str | None = None

class PromptRequestModel(BaseModel):
    prompt: str
    webhook_url: str | None = None

class SaveRequest(BaseModel):
    message: str | None = None

class DemoRequest(BaseModel):
    platform: str = "web"
    webhook_url: str | None = None

class BuildStatusResponse(BaseModel):
    build_id: str
    slug: str
    status: str
    phase: str
    message: str
    platforms: dict = {}
    elapsed_seconds: int = 0
    logs: list[str] = []
    warnings: list[str] = []
    capabilities: dict | None = None

class WorkspaceResponse(BaseModel):
    slug: str
    path: str
    platform: str
    owner_id: int | None = None
    account_id: str | None = None
    capabilities: dict | None = None

class RenameRequest(BaseModel):
    new_name: str

class PlatformBuildRequest(BaseModel):
    platform: str = "web"
    webhook_url: str | None = None

class PlanAppRequest(BaseModel):
    description: str

class CommitRequest(BaseModel):
    message: str | None = None
    auto_push: bool = False

class BranchRequest(BaseModel):
    name: str | None = None

class StashRequest(BaseModel):
    pop: bool = False

class GitResponse(BaseModel):
    output: str

class SaveListResponse(BaseModel):
    saves: list[dict] = []
    message: str

class RegisterRequest(BaseModel):
    display_name: str
    email: str | None = None

class CredentialRequest(BaseModel):
    """Generic credential payload — contents depend on type."""
    data: dict

class CreateKeyRequest(BaseModel):
    label: str = "default"

class WhitelistRequest(BaseModel):
    account_id: str

class ExtractRequest(BaseModel):
    """LLM-powered structured extraction from a document."""
    text: str
    json_schema: dict
    provider: str | None = None   # override auto-detection (anthropic/openai/groq/google/...)
    model: str | None = None      # override provider default
    system_prompt: str | None = None
    temperature: float = 0.1


class ExtractDocTextRequest(BaseModel):
    """Server-side text extraction from a .pdf or .docx file."""
    filename: str
    base64: str


class GeocodeVenueInput(BaseModel):
    """Single venue to resolve via geocoding."""
    id: str
    name: str
    address: str | None = None


class GeocodeVenuesRequest(BaseModel):
    """Batch geocoding request — one round-trip per import flow."""
    venues: list[GeocodeVenueInput]


class GolfCourseLookupInput(BaseModel):
    """One course-name query for /api/v1/golf-course-lookup."""
    id: str                       # caller-supplied (e.g. the venueId we'll bind to)
    name: str                     # "Revere Golf Club — Concord course"
    city: str | None = None       # optional disambiguator


class GolfCourseLookupRequest(BaseModel):
    """Batch golf-course lookup."""
    courses: list[GolfCourseLookupInput]


# ── Helpers ──────────────────────────────────────────────────────────────────

def _require_admin(account: Account):
    if account.role != "admin":
        raise HTTPException(403, "Admin access required")


def _check_workspace_access(account: Account, slug: str):
    """Ensure the account can access this workspace."""
    if account.role == "admin":
        return
    registry = service._get_registry()
    if not registry.can_access(slug, user_id=0, is_admin=False, account_id=account.account_id):
        raise HTTPException(403, f"No access to workspace '{slug}'")


def _status_to_response(s: service.BuildStatus) -> BuildStatusResponse:
    return BuildStatusResponse(
        build_id=s.build_id,
        slug=s.slug,
        status=s.status,
        phase=s.phase,
        message=s.message,
        platforms=s.platforms,
        elapsed_seconds=s.elapsed_seconds,
        logs=s.logs,
        warnings=getattr(s, 'warnings', []),
        capabilities=getattr(s, 'capabilities', None),
    )


# ── Registration (no auth required) ─────────────────────────────────────────

@app.post("/api/v1/register")
async def register(req: RegisterRequest):
    """Register a new account. Returns account_id and API key (shown once)."""
    mgr = _get_account_mgr()
    acct, raw_key = mgr.register(req.display_name, email=req.email)
    return {
        "account_id": acct.account_id,
        "display_name": acct.display_name,
        "api_key": raw_key,
        "message": "Save your API key — it cannot be retrieved later.",
    }


# ── Account endpoints (authenticated) ───────────────────────────────────────

@app.get("/api/v1/account")
async def get_account(account: Account = Depends(get_current_account)):
    """Get current account info, capabilities, and setup checklist."""
    mgr = _get_account_mgr()
    return {
        "account_id": account.account_id,
        "display_name": account.display_name,
        "email": account.email,
        "role": account.role,
        "discord_user_id": account.discord_user_id,
        "shared_store_access": account.shared_store_access,
        "capabilities": mgr.get_capabilities(account.account_id),
        "agent_provider": get_provider_capabilities().__dict__,
        "setup_checklist": mgr.get_setup_checklist(account.account_id),
        "created_at": account.created_at,
    }


@app.post("/api/v1/account/credentials/{cred_type}")
async def set_credential(cred_type: str, req: CredentialRequest,
                         account: Account = Depends(get_current_account)):
    """Set a credential (llm, supabase, apple, google)."""
    mgr = _get_account_mgr()
    ok = mgr.set_credential(account.account_id, cred_type, req.data)
    if not ok:
        raise HTTPException(400, f"Invalid credential type: {cred_type}")
    return {
        "credential_type": cred_type,
        "status": "stored",
        "capabilities": mgr.get_capabilities(account.account_id),
    }


@app.get("/api/v1/account/credentials")
async def list_credentials(account: Account = Depends(get_current_account)):
    """List which credential types are configured (no secrets exposed)."""
    mgr = _get_account_mgr()
    return mgr.list_credentials(account.account_id)


@app.delete("/api/v1/account/credentials/{cred_type}")
async def delete_credential(cred_type: str, account: Account = Depends(get_current_account)):
    """Remove a credential."""
    mgr = _get_account_mgr()
    ok = mgr.delete_credential(account.account_id, cred_type)
    if not ok:
        raise HTTPException(404, f"Credential '{cred_type}' not found")
    return {
        "credential_type": cred_type,
        "status": "deleted",
        "capabilities": mgr.get_capabilities(account.account_id),
    }


@app.post("/api/v1/extract")
async def extract_structured(req: ExtractRequest,
                             account: Account = Depends(get_current_account)):
    """
    Extract structured JSON from a document using the caller's configured LLM key.

    Auth: standard Bearer token (same as other endpoints).
    LLM key: read from the account's `llm` credential (set via
             POST /api/v1/account/credentials/llm). Provider is auto-detected
             from the key prefix unless `provider` is passed explicitly.

    Response on success: {"data": <obj matching json_schema>, "provider": "...", "model": "..."}
    Response on error:   {"error": true, "error_message": "..."}
    """
    import extract_limits
    from llm_providers import extract_json, detect_provider, list_providers

    is_admin = account.role == "admin"
    allowed, used, limit = extract_limits.check_and_consume(
        account.account_id, is_admin=is_admin,
    )
    if not allowed:
        raise HTTPException(
            429,
            f"Daily extract limit reached ({used}/{limit}). Try again tomorrow.",
        )

    mgr = _get_account_mgr()
    cred = mgr.get_credential(account.account_id, "llm")
    if not cred or not cred.get("api_key"):
        raise HTTPException(
            400,
            "No LLM credential set. POST /api/v1/account/credentials/llm with "
            '{"data": {"api_key": "sk-..."}} first.',
        )
    api_key = cred["api_key"]

    provider = req.provider or detect_provider(api_key)
    if not provider:
        raise HTTPException(
            400,
            f"Could not auto-detect provider from key prefix. Pass 'provider' explicitly. "
            f"Supported: {list_providers()}",
        )

    result = await extract_json(
        api_key=api_key,
        text=req.text,
        json_schema=req.json_schema,
        provider=provider,
        model=req.model,
        system_prompt=req.system_prompt,
        temperature=req.temperature,
    )
    if result.get("error"):
        # Return 200 with error body so callers can surface model errors without
        # HTTP-level retries kicking in. Swap to raise if you prefer HTTP errors.
        return result
    return result


@app.post("/api/v1/extract-doc-text")
async def extract_doc_text(req: ExtractDocTextRequest,
                           account: Account = Depends(get_current_account)):
    """
    Extract plain text from a .pdf / .docx file. Called by mobile clients that
    collected the file via a platform file picker. Client-side .txt/.md uploads
    should skip this endpoint.

    Request: {"filename": "<name.ext>", "base64": "<bytes>"}
    Response: {"text": "<extracted>", "pages": <int>} or HTTPException on failure.
    """
    import extract_limits

    is_admin = account.role == "admin"
    allowed, used, limit = extract_limits.check_and_consume(
        account.account_id, is_admin=is_admin,
    )
    if not allowed:
        raise HTTPException(
            429,
            f"Daily extract limit reached ({used}/{limit}). Try again tomorrow.",
        )

    import base64
    import io

    ext = req.filename.rsplit(".", 1)[-1].lower() if "." in req.filename else ""
    if ext not in {"pdf", "docx"}:
        raise HTTPException(400, f"Unsupported file type '.{ext}'. Use .pdf or .docx.")

    try:
        raw = base64.b64decode(req.base64, validate=False)
    except Exception as e:
        raise HTTPException(400, f"Invalid base64: {e}")

    if len(raw) == 0:
        raise HTTPException(400, "Empty file.")
    if len(raw) > 25 * 1024 * 1024:
        raise HTTPException(413, "File too large (>25 MB).")

    if ext == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise HTTPException(500, "pypdf not installed on the bridge.")
        try:
            reader = PdfReader(io.BytesIO(raw))
            pages = [(p.extract_text() or "") for p in reader.pages]
            text = "\n\n".join(pages).strip()
            return {"text": text, "pages": len(pages)}
        except Exception as e:
            raise HTTPException(422, f"PDF parse failed: {e}")

    # .docx
    try:
        import docx  # python-docx
    except ImportError:
        raise HTTPException(500, "python-docx not installed on the bridge.")
    try:
        document = docx.Document(io.BytesIO(raw))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts).strip()
        return {"text": text, "pages": len(document.paragraphs)}
    except Exception as e:
        raise HTTPException(422, f"DOCX parse failed: {e}")


@app.post("/api/v1/geocode-venues")
async def geocode_venues(req: GeocodeVenuesRequest,
                         account: Account = Depends(get_current_account)):
    """
    Batch-resolve venue name+address to lat/lng + canonical address via Mapbox.

    Per-venue result has `status`:
      "ok"         → latitude, longitude, canonical_address, place_id set
      "not_found"  → Mapbox returned no features for the query
      "error"      → transient failure; message in `message`
      "skipped"    → empty input (nothing to geocode)

    Caller keeps the returned `id` to merge results back into its venue list.
    Admin-paid for now (uses bridge-side MAPBOX_TOKEN); will move to BYOK later.
    """
    import httpx

    mapbox_token = os.getenv("MAPBOX_TOKEN", "").strip()
    if not mapbox_token:
        raise HTTPException(
            500,
            "MAPBOX_TOKEN not configured on bridge. Set it in .env and restart pm2.",
        )

    results: list[dict] = []
    async with httpx.AsyncClient(timeout=10.0) as client:
        for v in req.venues:
            query = ", ".join(p for p in (v.name.strip(), (v.address or "").strip()) if p)
            if not query:
                results.append({"id": v.id, "status": "skipped"})
                continue
            try:
                resp = await client.get(
                    "https://api.mapbox.com/search/geocode/v6/forward",
                    params={
                        "q": query,
                        "access_token": mapbox_token,
                        "limit": 1,
                    },
                )
                resp.raise_for_status()
                data = resp.json()
                feats = data.get("features") or []
                if not feats:
                    results.append({"id": v.id, "status": "not_found", "query": query})
                    continue
                f = feats[0]
                coords = (f.get("geometry") or {}).get("coordinates") or [None, None]
                props = f.get("properties") or {}
                results.append({
                    "id": v.id,
                    "status": "ok",
                    "latitude": coords[1],
                    "longitude": coords[0],
                    "canonical_address": props.get("full_address") or props.get("place_formatted"),
                    "place_id": props.get("mapbox_id"),
                })
            except httpx.HTTPStatusError as e:
                results.append({
                    "id": v.id,
                    "status": "error",
                    "message": f"Mapbox {e.response.status_code}: {e.response.text[:200]}",
                })
            except Exception as e:
                results.append({"id": v.id, "status": "error", "message": str(e)[:200]})
    return {"results": results}


@app.post("/api/v1/golf-course-lookup")
async def golf_course_lookup(req: GolfCourseLookupRequest,
                             account: Account = Depends(get_current_account)):
    """
    Resolve a list of course names to canonical golfcourseapi.com data:
    canonical name, lat/lng, address, and per-tee par + handicap-stroke-index
    arrays for all 18 holes.

    Per-course result has `status`:
      "ok"        → canonical_name, latitude, longitude, address, tees[]
      "ambiguous" → multiple plausible matches; `candidates[]` lists top 3
      "not_found" → API returned no matches
      "error"     → transient failure; message in `message`

    Each tee entry: {tee_name, course_rating, slope, par_total, total_yards,
                     pars: [18], stroke_indices: [18]}.

    Admin-paid via bridge GOLF_COURSE_API_KEY env var; will move to BYOK when
    user-side credentials grow that far.
    """
    import httpx

    api_key = os.getenv("GOLF_COURSE_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(
            500,
            "GOLF_COURSE_API_KEY not configured on bridge. Set in .env, restart pm2.",
        )

    headers = {"Authorization": f"Key {api_key}", "Accept": "application/json"}

    def _normalize_tee(tee: dict) -> dict:
        holes = tee.get("holes") or []
        return {
            "tee_name": tee.get("tee_name"),
            "course_rating": tee.get("course_rating"),
            "slope": tee.get("slope_rating"),
            "par_total": tee.get("par_total"),
            "total_yards": tee.get("total_yards"),
            "number_of_holes": tee.get("number_of_holes"),
            "pars": [h.get("par") for h in holes],
            "stroke_indices": [h.get("handicap") for h in holes],
            "yardages": [h.get("yardage") for h in holes],
        }

    def _course_to_record(course: dict) -> dict:
        loc = course.get("location") or {}
        male = course.get("tees", {}).get("male") or []
        female = course.get("tees", {}).get("female") or []
        # Prefer male tees as the default set, fall back to female. Caller can
        # see all and let the user pick in review.
        all_tees = (male if male else female)
        # Canonical name: "<Club> — <Course>" except when the API returns
        # club_name == course_name (e.g. Reflection Bay), where the dash
        # variant produces "Reflection Bay Golf Club — Reflection Bay Golf
        # Club". Fall back to the club name alone in that case.
        club = (course.get("club_name") or "").strip()
        coursename = (course.get("course_name") or "").strip()
        if club and coursename and club.lower() != coursename.lower():
            canonical = f"{club} — {coursename}"
        else:
            canonical = club or coursename
        return {
            "course_id": course.get("id"),
            "canonical_name": canonical,
            "club_name": course.get("club_name"),
            "course_name": course.get("course_name"),
            "address": loc.get("address"),
            "latitude": loc.get("latitude"),
            "longitude": loc.get("longitude"),
            "tees": [_normalize_tee(t) for t in all_tees if t.get("holes")],
        }

    import asyncio, re

    def _core_query(raw_name: str) -> str:
        """Strip noise that confuses the search index.

        The free GolfCourseAPI search is a phrase match — it does not handle
        em-dashes, the literal word "course", or trailing city qualifiers
        well. We narrow to the club name plus any sub-course token.
        """
        s = raw_name.replace("—", " ").replace("–", " ").replace("-", " ")
        s = re.sub(r"\b(course|club|golf|country|the)\b", " ", s, flags=re.I)
        s = re.sub(r"\s+", " ", s).strip()
        return s or raw_name.strip()

    async def _search_with_retry(client, query: str, max_retries: int = 4):
        backoff = 1.5
        for attempt in range(max_retries):
            r = await client.get(
                "https://api.golfcourseapi.com/v1/search",
                params={"search_query": query},
            )
            if r.status_code == 429:
                await asyncio.sleep(backoff)
                backoff *= 2
                continue
            return r
        return r  # last attempt's response

    results: list[dict] = []
    async with httpx.AsyncClient(timeout=20.0, headers=headers) as client:
        for c in req.courses:
            raw = c.name.strip()
            queries = []
            core = _core_query(raw)
            queries.append(core)
            # Sub-course token (e.g. "Concord" inside "Revere — Concord course")
            sub = re.search(r"(?:—|–|-)\s*([A-Za-z]+)", raw)
            if sub:
                queries.append(f"{core} {sub.group(1)}")
            queries.append(raw)  # last-resort: the literal user string
            try:
                candidates = []
                last_query_used = queries[0]
                for q in queries:
                    last_query_used = q
                    search = await _search_with_retry(client, q)
                    if search.status_code != 200:
                        # Surface only on the final query attempt
                        if q == queries[-1]:
                            results.append({
                                "id": c.id, "status": "error",
                                "message": f"search HTTP {search.status_code}: {search.text[:160]}",
                            })
                            candidates = "_FAILED_"
                            break
                        else:
                            continue
                    candidates = (search.json() or {}).get("courses") or []
                    if candidates:
                        break
                if candidates == "_FAILED_":
                    continue
                if not candidates:
                    results.append({
                        "id": c.id, "status": "not_found", "query": last_query_used,
                    })
                    continue
                # Throttle — free tier is harsh.
                await asyncio.sleep(0.6)

                # 2. Strict-match on the user-supplied name. If the query
                # mentions both a club ("Revere") and a course ("Concord"),
                # prefer the candidate whose course_name matches.
                lower_name = c.name.lower()
                def _score(course: dict) -> int:
                    club = (course.get("club_name") or "").lower()
                    cn   = (course.get("course_name") or "").lower()
                    s = 0
                    for tok in lower_name.replace("—", " ").replace("-", " ").split():
                        if tok in club: s += 1
                        if tok in cn:   s += 2  # course_name match outranks club match
                    return s

                ranked = sorted(candidates, key=_score, reverse=True)
                best = ranked[0]
                top_score = _score(best)
                # If runner-up scores within 1, surface as ambiguous so the user
                # picks in review.
                runner_up_score = _score(ranked[1]) if len(ranked) > 1 else -1

                if top_score == 0:
                    results.append({
                        "id": c.id, "status": "not_found", "query": query,
                        "candidates": [
                            {"course_id": ck.get("id"),
                             "canonical_name": f"{ck.get('club_name')} — {ck.get('course_name')}".strip(" —"),
                             "city": (ck.get("location") or {}).get("city"),
                             "state": (ck.get("location") or {}).get("state")}
                            for ck in ranked[:3]
                        ],
                    })
                    continue

                # 3. Fetch full detail (search payload sometimes truncates tees).
                # Throttle the detail call too.
                await asyncio.sleep(0.6)
                detail = None
                for attempt in range(4):
                    detail = await client.get(
                        f"https://api.golfcourseapi.com/v1/courses/{best['id']}"
                    )
                    if detail.status_code != 429:
                        break
                    await asyncio.sleep(1.5 * (2 ** attempt))
                if detail.status_code != 200:
                    # Fall back to whatever the search returned.
                    rec = _course_to_record(best)
                else:
                    rec = _course_to_record((detail.json() or {}).get("course") or best)

                rec["id"] = c.id
                rec["status"] = "ambiguous" if runner_up_score >= top_score - 1 and len(ranked) > 1 else "ok"
                if rec["status"] == "ambiguous":
                    rec["candidates"] = [
                        {"course_id": ck.get("id"),
                         "canonical_name": f"{ck.get('club_name')} — {ck.get('course_name')}".strip(" —"),
                         "city": (ck.get("location") or {}).get("city"),
                         "state": (ck.get("location") or {}).get("state")}
                        for ck in ranked[:3]
                    ]
                results.append(rec)

            except httpx.HTTPStatusError as e:
                results.append({
                    "id": c.id, "status": "error",
                    "message": f"golfcourseapi {e.response.status_code}: {e.response.text[:200]}",
                })
            except Exception as e:
                results.append({"id": c.id, "status": "error", "message": str(e)[:200]})

    return {"results": results}


@app.post("/api/v1/account/keys")
async def create_api_key(req: CreateKeyRequest = CreateKeyRequest(),
                         account: Account = Depends(get_current_account)):
    """Create a new API key for this account."""
    mgr = _get_account_mgr()
    raw_key = mgr.create_api_key(account.account_id, req.label)
    if not raw_key:
        raise HTTPException(500, "Failed to create API key")
    return {
        "api_key": raw_key,
        "label": req.label,
        "message": "Save your API key — it cannot be retrieved later.",
    }


@app.get("/api/v1/account/keys")
async def list_api_keys(account: Account = Depends(get_current_account)):
    """List API keys (prefix + label only, never the full key)."""
    mgr = _get_account_mgr()
    return mgr.list_api_keys(account.account_id)


@app.delete("/api/v1/account/keys/{prefix}")
async def revoke_api_key(prefix: str, account: Account = Depends(get_current_account)):
    """Revoke an API key by its prefix."""
    mgr = _get_account_mgr()
    ok = mgr.revoke_api_key(account.account_id, prefix)
    if not ok:
        raise HTTPException(404, f"API key with prefix '{prefix}' not found")
    return {"prefix": prefix, "status": "revoked"}


# ── Admin endpoints ─────────────────────────────────────────────────────────

@app.post("/api/v1/admin/whitelist")
async def grant_shared_store(req: WhitelistRequest,
                             account: Account = Depends(get_current_account)):
    """Grant shared store access to an account (admin only)."""
    _require_admin(account)
    mgr = _get_account_mgr()
    ok = mgr.set_shared_store_access(req.account_id, True)
    if not ok:
        raise HTTPException(404, f"Account '{req.account_id}' not found")
    return {"account_id": req.account_id, "shared_store_access": True}


@app.delete("/api/v1/admin/whitelist/{account_id}")
async def revoke_shared_store(account_id: str,
                              account: Account = Depends(get_current_account)):
    """Revoke shared store access (admin only)."""
    _require_admin(account)
    mgr = _get_account_mgr()
    ok = mgr.set_shared_store_access(account_id, False)
    if not ok:
        raise HTTPException(404, f"Account '{account_id}' not found")
    return {"account_id": account_id, "shared_store_access": False}


@app.get("/api/v1/admin/whitelist")
async def list_whitelist(account: Account = Depends(get_current_account)):
    """List accounts with shared store access (admin only)."""
    _require_admin(account)
    mgr = _get_account_mgr()
    return [
        {"account_id": a.account_id, "display_name": a.display_name}
        for a in mgr.list_accounts() if a.shared_store_access
    ]


# ── Existing endpoints (updated for multi-tenant auth) ──────────────────────

@app.post("/api/v1/buildapp", response_model=BuildStatusResponse)
async def buildapp(req: BuildAppRequest, account: Account = Depends(get_current_account)):
    """Start a new app build. Returns immediately with a build_id for polling."""
    build_req = service.BuildRequest(
        description=req.description,
        app_name=req.app_name,
        platform=req.platform,
        skip_supabase=req.skip_supabase,
        webhook_url=req.webhook_url,
        account_id=account.account_id,
    )
    status = await service.build_app(build_req)
    return _status_to_response(status)


@app.get("/api/v1/builds/{build_id}", response_model=BuildStatusResponse)
async def get_build(build_id: str, account: Account = Depends(get_current_account)):
    """Poll build status."""
    status = service.get_build(build_id)
    if not status:
        raise HTTPException(404, f"Build '{build_id}' not found")
    return _status_to_response(status)


@app.get("/api/v1/workspaces", response_model=list[WorkspaceResponse])
async def list_workspaces(account: Account = Depends(get_current_account)):
    workspaces = await service.list_workspaces(account_id=account.account_id)
    mgr = _get_account_mgr()
    caps = mgr.get_capabilities(account.account_id)
    return [
        WorkspaceResponse(
            slug=w.slug, path=w.path, platform=w.platform,
            owner_id=w.owner_id, account_id=w.account_id,
            capabilities=caps,
        )
        for w in workspaces
    ]


@app.get("/api/v1/workspaces/{slug}", response_model=WorkspaceResponse)
async def get_workspace(slug: str, account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    ws = await service.get_workspace(slug)
    if not ws:
        raise HTTPException(404, f"Workspace '{slug}' not found")
    mgr = _get_account_mgr()
    return WorkspaceResponse(
        slug=ws.slug, path=ws.path, platform=ws.platform,
        owner_id=ws.owner_id, account_id=ws.account_id,
        capabilities=mgr.get_capabilities(account.account_id),
    )


@app.post("/api/v1/workspaces/{slug}/prompt", response_model=BuildStatusResponse)
async def send_prompt(slug: str, req: PromptRequestModel,
                      account: Account = Depends(get_current_account)):
    """Send a prompt to a workspace. Returns build_id for polling."""
    _check_workspace_access(account, slug)
    try:
        status = await service.send_prompt(
            service.PromptRequest(workspace=slug, prompt=req.prompt, webhook_url=req.webhook_url)
        )
        return _status_to_response(status)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/workspaces/{slug}/demo", response_model=BuildStatusResponse)
async def demo(slug: str, req: DemoRequest = DemoRequest(),
               account: Account = Depends(get_current_account)):
    """Trigger a demo build for a workspace."""
    _check_workspace_access(account, slug)
    try:
        status = await service.demo_workspace(slug, req.platform)
        status.webhook_url = req.webhook_url
        return _status_to_response(status)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/workspaces/{slug}/save")
async def save(slug: str, req: SaveRequest = SaveRequest(),
               account: Account = Depends(get_current_account)):
    """Save/checkpoint a workspace."""
    _check_workspace_access(account, slug)
    try:
        return await service.save_workspace(slug, req.message)
    except ValueError as e:
        raise HTTPException(404, str(e))


# ── Workspace Management ────────────────────────────────────────────────────

@app.post("/api/v1/workspaces/{slug}/use")
async def set_active_workspace(slug: str, account: Account = Depends(get_current_account)):
    """Set workspace as the active default."""
    _check_workspace_access(account, slug)
    ok = await service.set_default_workspace(account.account_id, slug)
    if not ok:
        raise HTTPException(404, f"Workspace '{slug}' not found")
    return {"slug": slug, "active": True}


@app.patch("/api/v1/workspaces/{slug}")
async def rename_workspace(slug: str, req: RenameRequest,
                           account: Account = Depends(get_current_account)):
    """Rename a workspace."""
    _check_workspace_access(account, slug)
    ok = await service.rename_workspace(slug, req.new_name)
    if not ok:
        raise HTTPException(400, f"Rename failed — '{slug}' not found or '{req.new_name}' already exists")
    return {"old_slug": slug, "new_slug": req.new_name}


@app.delete("/api/v1/workspaces/{slug}")
async def delete_workspace(slug: str, force: bool = False,
                           account: Account = Depends(get_current_account)):
    """Delete a workspace (registry + files). Protected workspaces require force=true."""
    _check_workspace_access(account, slug)
    try:
        ok = await service.delete_workspace(slug, force=force)
    except ValueError as e:
        raise HTTPException(403, str(e))
    if not ok:
        raise HTTPException(404, f"Workspace '{slug}' not found")
    return {"slug": slug, "deleted": True}


@app.post("/api/v1/workspaces/{slug}/protect")
async def protect_workspace(slug: str, protected: bool = True,
                            account: Account = Depends(get_current_account)):
    """Toggle protection on a workspace."""
    _check_workspace_access(account, slug)
    registry = service._get_registry()
    if not registry.exists(slug):
        raise HTTPException(404, f"Workspace '{slug}' not found")
    registry.set_protected(slug, protected)
    return {"slug": slug, "protected": protected}


@app.post("/api/v1/workspaces/{slug}/newsession")
async def new_session(slug: str, account: Account = Depends(get_current_account)):
    """Clear Claude session for a workspace (start fresh conversation)."""
    _check_workspace_access(account, slug)
    ws = await service.get_workspace(slug)
    if not ws:
        raise HTTPException(404, f"Workspace '{slug}' not found")
    await service.clear_session(slug)
    return {"slug": slug, "session_cleared": True}


# ── Build & Plan ────────────────────────────────────────────────────────────

@app.post("/api/v1/workspaces/{slug}/build", response_model=BuildStatusResponse)
async def build_workspace(slug: str, req: PlatformBuildRequest = PlatformBuildRequest(),
                          account: Account = Depends(get_current_account)):
    """Build workspace for a specific platform. Returns build_id for polling."""
    _check_workspace_access(account, slug)
    try:
        status = await service.build_workspace_platform(slug, req.platform)
        status.webhook_url = req.webhook_url
        return _status_to_response(status)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/planapp")
async def plan_app(req: PlanAppRequest, account: Account = Depends(get_current_account)):
    """Generate an app plan from a description. Synchronous — returns plan JSON."""
    plan = await service.plan_app(req.description)
    if not plan:
        raise HTTPException(500, "Failed to generate plan")
    return plan


@app.post("/api/v1/workspaces/{slug}/appraise", response_model=BuildStatusResponse)
async def appraise_workspace(slug: str, account: Account = Depends(get_current_account)):
    """Run quality appraisal on a workspace. Returns build_id for polling."""
    _check_workspace_access(account, slug)
    try:
        status = await service.appraise_workspace(slug)
        return _status_to_response(status)
    except ValueError as e:
        raise HTTPException(404, str(e))


# ── Git Operations ──────────────────────────────────────────────────────────

@app.get("/api/v1/workspaces/{slug}/git/status", response_model=GitResponse)
async def git_status(slug: str, account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.git_status(slug)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.get("/api/v1/workspaces/{slug}/git/diff", response_model=GitResponse)
async def git_diff(slug: str, full: bool = False,
                   account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.git_diff(slug, full=full)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.get("/api/v1/workspaces/{slug}/git/log", response_model=GitResponse)
async def git_log(slug: str, count: int = 10,
                  account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.git_log(slug, count=count)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/workspaces/{slug}/git/commit", response_model=GitResponse)
async def git_commit(slug: str, req: CommitRequest = CommitRequest(),
                     account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.git_commit(slug, message=req.message, auto_push=req.auto_push)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/workspaces/{slug}/git/undo", response_model=GitResponse)
async def git_undo(slug: str, account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.git_undo(slug)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/workspaces/{slug}/git/branch", response_model=GitResponse)
async def git_branch(slug: str, req: BranchRequest = BranchRequest(),
                     account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.git_branch(slug, name=req.name)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/workspaces/{slug}/git/stash", response_model=GitResponse)
async def git_stash(slug: str, req: StashRequest = StashRequest(),
                    account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.git_stash(slug, pop=req.pop)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


# ── Save System ─────────────────────────────────────────────────────────────

@app.get("/api/v1/workspaces/{slug}/saves", response_model=SaveListResponse)
async def list_saves(slug: str, account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        message, saves = await service.save_list(slug)
        return SaveListResponse(saves=saves, message=message)
    except ValueError as e:
        raise HTTPException(404, str(e))


@app.post("/api/v1/workspaces/{slug}/saves/undo", response_model=GitResponse)
async def undo_save(slug: str, account: Account = Depends(get_current_account)):
    _check_workspace_access(account, slug)
    try:
        output = await service.save_undo(slug)
        return GitResponse(output=output)
    except ValueError as e:
        raise HTTPException(404, str(e))


# ── Smoke Tests ─────────────────────────────────────────────────────────────

class SmokeTestRequest(BaseModel):
    scenario: str | None = None  # "counter", "map", "video", or None for all
    api_tests: bool = False       # Also run API endpoint checks

@app.post("/api/v1/smoketest")
async def run_smoketest_endpoint(req: SmokeTestRequest = SmokeTestRequest(),
                                 account: Account = Depends(get_current_account)):
    """Kick off smoke tests. Returns a build_id for polling status."""
    _require_admin(account)
    import uuid
    import time as _time

    build_id = str(uuid.uuid4())[:8]
    status = BuildStatusResponse(
        build_id=build_id,
        slug="smoketest",
        status="running",
        phase="starting",
        message="Smoke test starting...",
    )
    _smoketest_results[build_id] = status

    async def _run():
        from helpers.smoketest_runner import run_smoketest, SCENARIO_NAMES
        from workspaces import WorkspaceRegistry
        from claude_runner import ClaudeRunner

        registry = WorkspaceRegistry()
        claude = ClaudeRunner()
        scenarios = [req.scenario] if req.scenario and req.scenario in SCENARIO_NAMES else None
        logs = []

        async def on_status(msg, file_path=None):
            cleaned = msg.replace("**", "").replace("`", "")
            logs.append(cleaned)
            status.logs = logs[-20:]  # keep last 20
            status.message = cleaned

        try:
            result = await run_smoketest(
                registry=registry,
                claude=claude,
                on_status=on_status,
                is_admin=False,
                scenarios=scenarios,
            )
            status.status = "success" if result.success else "failed"
            status.phase = "complete"
            status.message = result.summary()
            status.logs = logs

            # Run API tests if requested
            if req.api_tests:
                from helpers.api_smoketest import run_api_smoketest
                port = os.getenv("API_PORT", "8100")
                api_result = await run_api_smoketest(
                    base_url=f"http://localhost:{port}",
                    token=API_TOKEN,
                )
                status.message += "\n\n" + api_result.summary()
                if not api_result.success:
                    status.status = "failed"
        except Exception as e:
            logger.exception(f"[smoketest:{build_id}] Failed")
            status.status = "failed"
            status.phase = "complete"
            status.message = f"Error: {e}"

    import asyncio
    asyncio.create_task(_run())
    return status

@app.get("/api/v1/smoketest/{build_id}")
async def get_smoketest_status(build_id: str,
                               account: Account = Depends(get_current_account)):
    if build_id not in _smoketest_results:
        raise HTTPException(404, f"Smoke test '{build_id}' not found")
    return _smoketest_results[build_id]

_smoketest_results: dict[str, BuildStatusResponse] = {}


# ── Analytics ────────────────────────────────────────────────────────────────

@app.get("/api/v1/analytics")
async def analytics(account: Account = Depends(get_current_account)):
    """Build analytics: success rates, avg durations, per-workspace and per-operation breakdowns."""
    return service.get_analytics()


# ── Health ──────────────────────────────────────────────────────────────────

@app.get("/api/v1/health")
async def health():
    return {"status": "ok"}


# ── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    port = int(os.getenv("API_PORT", "8100"))
    logger.info(f"Starting API server on port {port}")
    logger.info(f"API token: {API_TOKEN[:8]}...")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
