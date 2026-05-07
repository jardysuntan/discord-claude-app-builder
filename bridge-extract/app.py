"""
app.py — Standalone FastAPI service exposing only the LLM-extract and
doc-text endpoints used by mobile clients (e.g. WereSoBach iOS).

Endpoints:
  GET  /healthz
  POST /api/v1/extract
  POST /api/v1/extract-doc-text

Auth: Bearer token, validated against the shared accounts.json (encrypted
credentials). Same scheme as the parent bridge so existing API keys and
admin BYOK creds keep working.

This service intentionally does NOT include any of the bridge's
shell-out / Claude CLI / build-runner code. It is safe to deploy to a
public host because it only routes LLM calls using caller-stored keys.
"""
import logging
import os
from typing import Optional

from fastapi import FastAPI, HTTPException, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from accounts import AccountManager, Account

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(name)s %(levelname)s %(message)s")
logger = logging.getLogger("bridge-extract")

# ── Account manager ─────────────────────────────────────────────────────────

_account_mgr: Optional[AccountManager] = None


def _get_account_mgr() -> AccountManager:
    global _account_mgr
    if _account_mgr is None:
        _account_mgr = AccountManager()
    return _account_mgr


async def get_current_account(authorization: str = Header(...)) -> Account:
    if not authorization.startswith("Bearer "):
        raise HTTPException(401, "Missing Bearer token")
    token = authorization[7:]
    acct = _get_account_mgr().authenticate(token)
    if not acct:
        raise HTTPException(403, "Invalid token")
    return acct


# ── App ─────────────────────────────────────────────────────────────────────

app = FastAPI(title="WereSoBach Bridge (extract)", version="1.0.0", docs_url="/api/docs")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/healthz")
async def healthz():
    return {"ok": True, "service": "bridge-extract"}


# ── Request models ──────────────────────────────────────────────────────────

class ExtractRequest(BaseModel):
    text: str
    json_schema: dict
    provider: str | None = None
    model: str | None = None
    system_prompt: str | None = None
    temperature: float = 0.1


class ExtractDocTextRequest(BaseModel):
    filename: str
    base64: str


class GeocodeVenueInput(BaseModel):
    id: str
    name: str
    address: str | None = None


class GeocodeVenuesRequest(BaseModel):
    venues: list[GeocodeVenueInput]
    # Optional trip-level location ("Las Vegas, NV", "Lexington, MA"). When
    # set, gets appended to every venue query so Mapbox's global ranking
    # doesn't pick the wrong international branch (e.g. Din Tai Fung in
    # Thailand for a Vegas trip).
    trip_location: str | None = None


class GeocodePlaceRequest(BaseModel):
    """Single-place autocomplete query for the trip-creation Location field.

    Distinct from /geocode-venues (which is a batch resolve for already-named
    POIs). This one returns up to N suggestions for a free-text query so the
    user can pick the right city. Limit kept small — autocomplete sends one
    request per debounce tick.
    """

    query: str
    limit: int = 5


class GolfCourseLookupInput(BaseModel):
    id: str
    name: str
    city: str | None = None


class GolfCourseLookupRequest(BaseModel):
    courses: list[GolfCourseLookupInput]


class ClassifyHousingRequest(BaseModel):
    text: str
    provider: str | None = None
    model: str | None = None


# ── /api/v1/extract ─────────────────────────────────────────────────────────

@app.post("/api/v1/extract")
async def extract_structured(req: ExtractRequest,
                             account: Account = Depends(get_current_account)):
    import extract_limits
    from llm_providers import extract_json, detect_provider, list_providers

    is_admin = account.role == "admin"
    allowed, used, limit = extract_limits.check_and_consume(
        account.account_id, is_admin=is_admin,
    )
    if not allowed:
        raise HTTPException(
            429,
            f"Daily extract limit reached ({used}/{limit}). Try again tomorrow.",
        )

    mgr = _get_account_mgr()
    cred = mgr.get_credential(account.account_id, "llm")
    if not cred or not cred.get("api_key"):
        raise HTTPException(
            400,
            "No LLM credential set. POST /api/v1/account/credentials/llm with "
            '{"data": {"api_key": "sk-..."}} on the parent bridge first.',
        )
    api_key = cred["api_key"]

    provider = req.provider or detect_provider(api_key)
    if not provider:
        raise HTTPException(
            400,
            f"Could not auto-detect provider from key prefix. Pass 'provider' explicitly. "
            f"Supported: {list_providers()}",
        )

    result = await extract_json(
        api_key=api_key,
        text=req.text,
        json_schema=req.json_schema,
        provider=provider,
        model=req.model,
        system_prompt=req.system_prompt,
        temperature=req.temperature,
    )
    return result


# ── /api/v1/extract-doc-text ────────────────────────────────────────────────

@app.post("/api/v1/extract-doc-text")
async def extract_doc_text(req: ExtractDocTextRequest,
                           account: Account = Depends(get_current_account)):
    import extract_limits

    is_admin = account.role == "admin"
    allowed, used, limit = extract_limits.check_and_consume(
        account.account_id, is_admin=is_admin,
    )
    if not allowed:
        raise HTTPException(
            429,
            f"Daily extract limit reached ({used}/{limit}). Try again tomorrow.",
        )

    import base64
    import io

    ext = req.filename.rsplit(".", 1)[-1].lower() if "." in req.filename else ""
    if ext not in {"pdf", "docx"}:
        raise HTTPException(400, f"Unsupported file type '.{ext}'. Use .pdf or .docx.")

    try:
        raw = base64.b64decode(req.base64, validate=False)
    except Exception as e:
        raise HTTPException(400, f"Invalid base64: {e}")

    if len(raw) == 0:
        raise HTTPException(400, "Empty file.")
    if len(raw) > 25 * 1024 * 1024:
        raise HTTPException(413, "File too large (>25 MB).")

    if ext == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise HTTPException(500, "pypdf not installed.")
        try:
            reader = PdfReader(io.BytesIO(raw))
            pages = [(p.extract_text() or "") for p in reader.pages]
            text = "\n\n".join(pages).strip()
            return {"text": text, "pages": len(pages)}
        except Exception as e:
            raise HTTPException(422, f"PDF parse failed: {e}")

    # .docx
    try:
        import docx  # python-docx
    except ImportError:
        raise HTTPException(500, "python-docx not installed.")
    try:
        document = docx.Document(io.BytesIO(raw))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts).strip()
        return {"text": text, "pages": len(document.paragraphs)}
    except Exception as e:
        raise HTTPException(422, f"DOCX parse failed: {e}")


# ── /api/v1/geocode-venues (Mapbox passthrough) ─────────────────────────────

@app.post("/api/v1/geocode-venues")
async def geocode_venues(req: GeocodeVenuesRequest,
                         account: Account = Depends(get_current_account)):
    import httpx

    mapbox_token = os.getenv("MAPBOX_TOKEN", "").strip()
    if not mapbox_token:
        raise HTTPException(500, "MAPBOX_TOKEN not configured.")

    results: list[dict] = []
    trip_loc = (req.trip_location or "").strip()
    async with httpx.AsyncClient(timeout=10.0) as client:
        for v in req.venues:
            parts = [v.name.strip(), (v.address or "").strip()]
            # If the venue extraction had no address, append the trip's own
            # location as a region hint. Skip the append when the venue's
            # address already names a city (avoids "Las Vegas, NV, Las Vegas, NV").
            has_city_hint = any(
                trip_loc and trip_loc.lower().split(",")[0].strip() in p.lower()
                for p in parts if p
            )
            if trip_loc and not has_city_hint:
                parts.append(trip_loc)
            query = ", ".join(p for p in parts if p)
            if not query:
                results.append({"id": v.id, "status": "skipped"})
                continue
            try:
                resp = await client.get(
                    "https://api.mapbox.com/search/geocode/v6/forward",
                    params={"q": query, "access_token": mapbox_token, "limit": 1},
                )
                resp.raise_for_status()
                data = resp.json()
                feats = data.get("features") or []
                if not feats:
                    results.append({"id": v.id, "status": "not_found", "query": query})
                    continue
                f = feats[0]
                coords = (f.get("geometry") or {}).get("coordinates") or [None, None]
                props = f.get("properties") or {}
                results.append({
                    "id": v.id,
                    "status": "ok",
                    "latitude": coords[1],
                    "longitude": coords[0],
                    "canonical_address": props.get("full_address") or props.get("place_formatted"),
                    "place_id": props.get("mapbox_id"),
                })
            except httpx.HTTPStatusError as e:
                results.append({
                    "id": v.id, "status": "error",
                    "message": f"Mapbox {e.response.status_code}: {e.response.text[:200]}",
                })
            except Exception as e:
                results.append({"id": v.id, "status": "error", "message": str(e)[:200]})
    return {"results": results}


# ── /api/v1/geocode-place (Mapbox autocomplete passthrough) ─────────────────

@app.post("/api/v1/geocode-place")
async def geocode_place(req: GeocodePlaceRequest,
                        account: Account = Depends(get_current_account)):
    """Forward-geocode a single user-typed query to a small list of city/region
    suggestions. Used by the iOS app's trip-creation Location field for
    autocomplete. Skipped for queries shorter than 2 chars to keep Mapbox
    request volume sane during fast typing."""
    import httpx

    mapbox_token = os.getenv("MAPBOX_TOKEN", "").strip()
    if not mapbox_token:
        raise HTTPException(500, "MAPBOX_TOKEN not configured.")

    query = (req.query or "").strip()
    if len(query) < 2:
        return {"suggestions": []}

    limit = max(1, min(req.limit, 10))

    async with httpx.AsyncClient(timeout=8.0) as client:
        try:
            resp = await client.get(
                "https://api.mapbox.com/search/geocode/v6/forward",
                params={
                    "q": query,
                    "access_token": mapbox_token,
                    "limit": limit,
                    # City-ish granularity. Skip POI/address (that's what
                    # /geocode-venues is for).
                    "types": "place,locality,region,district,country",
                },
            )
            resp.raise_for_status()
            data = resp.json()
        except httpx.HTTPStatusError as e:
            raise HTTPException(
                502, f"Mapbox {e.response.status_code}: {e.response.text[:200]}",
            )
        except Exception as e:
            raise HTTPException(502, f"Mapbox geocode failed: {e}")

    suggestions = []
    for f in (data.get("features") or [])[:limit]:
        coords = (f.get("geometry") or {}).get("coordinates") or [None, None]
        props = f.get("properties") or {}
        suggestions.append({
            "name": props.get("name") or "",
            "full_address": props.get("full_address") or props.get("place_formatted") or "",
            "latitude": coords[1],
            "longitude": coords[0],
            "place_id": props.get("mapbox_id"),
        })
    return {"suggestions": suggestions}


# ── /api/v1/golf-course-lookup (golfcourseapi.com passthrough) ──────────────

def _normalize_tee(tee: dict) -> dict:
    holes = tee.get("holes") or []
    return {
        "tee_name": tee.get("tee_name"),
        "course_rating": tee.get("course_rating"),
        "slope": tee.get("slope_rating"),
        "par_total": tee.get("par_total"),
        "total_yards": tee.get("total_yards"),
        "number_of_holes": tee.get("number_of_holes"),
        "pars": [h.get("par") for h in holes],
        "stroke_indices": [h.get("handicap") for h in holes],
        "yardages": [h.get("yardage") for h in holes],
    }


def _course_to_record(course: dict) -> dict:
    loc = course.get("location") or {}
    male = course.get("tees", {}).get("male") or []
    female = course.get("tees", {}).get("female") or []
    all_tees = male if male else female
    club = (course.get("club_name") or "").strip()
    coursename = (course.get("course_name") or "").strip()
    if club and coursename and club.lower() != coursename.lower():
        canonical = f"{club} — {coursename}"
    else:
        canonical = club or coursename
    return {
        "course_id": course.get("id"),
        "canonical_name": canonical,
        "club_name": course.get("club_name"),
        "course_name": course.get("course_name"),
        "address": loc.get("address"),
        "latitude": loc.get("latitude"),
        "longitude": loc.get("longitude"),
        "tees": [_normalize_tee(t) for t in all_tees if t.get("holes")],
    }


@app.post("/api/v1/golf-course-lookup")
async def golf_course_lookup(req: GolfCourseLookupRequest,
                             account: Account = Depends(get_current_account)):
    import asyncio
    import re
    import httpx

    api_key = os.getenv("GOLF_COURSE_API_KEY", "").strip()
    if not api_key:
        raise HTTPException(500, "GOLF_COURSE_API_KEY not configured.")

    headers = {"Authorization": f"Key {api_key}", "Accept": "application/json"}

    def _core_query(raw_name: str) -> str:
        s = raw_name.replace("—", " ").replace("–", " ").replace("-", " ")
        s = re.sub(r"\b(course|club|golf|country|the)\b", " ", s, flags=re.I)
        s = re.sub(r"\s+", " ", s).strip()
        return s or raw_name.strip()

    async def _search_with_retry(client, query: str, max_retries: int = 4):
        backoff = 1.5
        r = None
        for _ in range(max_retries):
            r = await client.get(
                "https://api.golfcourseapi.com/v1/search",
                params={"search_query": query},
            )
            if r.status_code == 429:
                await asyncio.sleep(backoff)
                backoff *= 2
                continue
            return r
        return r

    results: list[dict] = []
    async with httpx.AsyncClient(timeout=20.0, headers=headers) as client:
        for c in req.courses:
            raw = c.name.strip()
            queries = [_core_query(raw)]
            sub = re.search(r"(?:—|–|-)\s*([A-Za-z]+)", raw)
            if sub:
                queries.append(f"{queries[0]} {sub.group(1)}")
            queries.append(raw)
            try:
                candidates = []
                last_query_used = queries[0]
                for q in queries:
                    last_query_used = q
                    search = await _search_with_retry(client, q)
                    if search.status_code != 200:
                        if q == queries[-1]:
                            results.append({
                                "id": c.id, "status": "error",
                                "message": f"search HTTP {search.status_code}: {search.text[:160]}",
                            })
                            candidates = "_FAILED_"
                            break
                        continue
                    candidates = (search.json() or {}).get("courses") or []
                    if candidates:
                        break
                if candidates == "_FAILED_":
                    continue
                if not candidates:
                    results.append({"id": c.id, "status": "not_found", "query": last_query_used})
                    continue
                await asyncio.sleep(0.6)

                lower_name = c.name.lower()

                def _score(course: dict) -> int:
                    club = (course.get("club_name") or "").lower()
                    cn = (course.get("course_name") or "").lower()
                    s = 0
                    for tok in lower_name.replace("—", " ").replace("-", " ").split():
                        if tok in club:
                            s += 1
                        if tok in cn:
                            s += 2
                    return s

                ranked = sorted(candidates, key=_score, reverse=True)
                best = ranked[0]
                top_score = _score(best)
                runner_up_score = _score(ranked[1]) if len(ranked) > 1 else -1

                if top_score == 0:
                    results.append({
                        "id": c.id, "status": "not_found", "query": last_query_used,
                        "candidates": [
                            {"course_id": ck.get("id"),
                             "canonical_name": f"{ck.get('club_name')} — {ck.get('course_name')}".strip(" —"),
                             "city": (ck.get("location") or {}).get("city"),
                             "state": (ck.get("location") or {}).get("state")}
                            for ck in ranked[:3]
                        ],
                    })
                    continue

                await asyncio.sleep(0.6)
                detail = None
                for attempt in range(4):
                    detail = await client.get(
                        f"https://api.golfcourseapi.com/v1/courses/{best['id']}"
                    )
                    if detail.status_code != 429:
                        break
                    await asyncio.sleep(1.5 * (2 ** attempt))
                if detail is None or detail.status_code != 200:
                    rec = _course_to_record(best)
                else:
                    rec = _course_to_record((detail.json() or {}).get("course") or best)

                rec["id"] = c.id
                rec["status"] = "ambiguous" if runner_up_score >= top_score - 1 and len(ranked) > 1 else "ok"
                if rec["status"] == "ambiguous":
                    rec["candidates"] = [
                        {"course_id": ck.get("id"),
                         "canonical_name": f"{ck.get('club_name')} — {ck.get('course_name')}".strip(" —"),
                         "city": (ck.get("location") or {}).get("city"),
                         "state": (ck.get("location") or {}).get("state")}
                        for ck in ranked[:3]
                    ]
                results.append(rec)

            except httpx.HTTPStatusError as e:
                results.append({
                    "id": c.id, "status": "error",
                    "message": f"golfcourseapi {e.response.status_code}: {e.response.text[:200]}",
                })
            except Exception as e:
                results.append({"id": c.id, "status": "error", "message": str(e)[:200]})

    return {"results": results}


# ── /api/v1/classify-housing (small one-shot LLM extract for paste-blob UX) ─

CLASSIFY_HOUSING_SCHEMA = {
    "type": "object",
    "properties": {
        "type": {
            "type": "string",
            "enum": ["hotel", "airbnb", "house", "rental", "other"],
            "description": "Best-guess accommodation type from the text.",
        },
        "name": {
            "type": "string",
            "description": "Listing name / property title if present (e.g. 'Modern 4BR Pool Villa'). Empty string if not found.",
        },
        "notes": {
            "type": "string",
            "description": "1-3 sentence summary of the listing description suitable as guest-facing notes. Skip boilerplate like host blurbs.",
        },
        "amenities": {
            "type": "array",
            "items": {"type": "string"},
            "description": "Short amenity labels like 'Pool', 'Hot tub', 'WiFi', 'Pet friendly'. Title case. Deduplicated.",
        },
        "checkIn": {
            "type": "string",
            "description": "ISO date YYYY-MM-DD if a check-in date is mentioned, else empty string.",
        },
        "checkOut": {
            "type": "string",
            "description": "ISO date YYYY-MM-DD if a check-out date is mentioned, else empty string.",
        },
    },
    "required": ["type", "name", "notes", "amenities", "checkIn", "checkOut"],
}

CLASSIFY_HOUSING_SYSTEM = (
    "You convert a pasted Airbnb / hotel / VRBO listing blurb into a "
    "compact structured housing record. Extract only what the text "
    "actually says. If a field isn't present, return an empty string "
    "(or empty array for amenities). Do not invent dates, addresses, "
    "or amenities."
)


@app.post("/api/v1/classify-housing")
async def classify_housing(req: ClassifyHousingRequest,
                           account: Account = Depends(get_current_account)):
    from llm_providers import extract_json, detect_provider, list_providers

    text = (req.text or "").strip()
    if not text:
        raise HTTPException(400, "Empty text.")
    if len(text) > 20_000:
        raise HTTPException(413, "Text too long (>20k chars).")

    mgr = _get_account_mgr()
    cred = mgr.get_credential(account.account_id, "llm")
    if not cred or not cred.get("api_key"):
        raise HTTPException(
            400,
            "No LLM credential set. POST /api/v1/account/credentials/llm with "
            '{"data": {"api_key": "sk-..."}} on the parent bridge first.',
        )
    api_key = cred["api_key"]

    provider = req.provider or detect_provider(api_key)
    if not provider:
        raise HTTPException(
            400,
            f"Could not auto-detect provider from key prefix. Pass 'provider' explicitly. "
            f"Supported: {list_providers()}",
        )

    return await extract_json(
        api_key=api_key,
        text=text,
        json_schema=CLASSIFY_HOUSING_SCHEMA,
        provider=provider,
        model=req.model,
        system_prompt=CLASSIFY_HOUSING_SYSTEM,
        temperature=0.1,
    )


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8100"))
    uvicorn.run(app, host="0.0.0.0", port=port)
